import os
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def extract_code(file_path):
    with open(file_path, "r", encoding="utf-8") as file:
        code = file.read()
    return code


def compile_and_run_main(main_file_path, arguments=[]):
    try:
        command = ["python", main_file_path] + arguments
        result = subprocess.run(
            command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
        )
        output = result.stdout + result.stderr
        return output
    except Exception as e:
        return str(e)


def get_main_arguments(executable):
    arguments = input(f"Введите параметры для выполнения {executable} (при наличии): ")
    return arguments.split()


def create_docx(directory, executable):
    if ".py" not in executable:
        executable = "main.py"

    output_file = os.path.join(directory, "report.docx")
    doc = Document()
    found_tests = False
    found_executable = False

    for style in doc.styles:
        if style.name == "Normal" or style.name.startswith("Heading"):
            style.font.name = "Times New Roman"
            style.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_heading("Текст программы", level=1)

    for root, dirs, files in os.walk(directory):
        if "__pycahce__" in root:
            continue
        for file in files:
            if file.endswith(".py") and "__init__" not in file:
                print(f"Обработан {file}...")
                file_path = os.path.join(root, file)

                code = extract_code(file_path)

                doc.add_heading(file, level=2)

                code_paragraph = doc.add_paragraph()
                code_paragraph.add_run(code).font.size = Pt(10)
                code_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for line in code.split("\n"):
                    code_paragraph.add_run(line).font.size = Pt(10)
                    code_paragraph.add_run("\n").font.size = Pt(10)

                if file == executable:
                    main_arguments = get_main_arguments(executable)
                    console_output = compile_and_run_main(file_path, main_arguments)
                    found_executable = True
                if file == "tests.py":
                    console_test_output = compile_and_run_main(file_path)
                    found_tests = True
    if found_executable:
        doc.add_heading("Результат выполнения программы", level=2)
        doc.add_paragraph(console_output)
    if found_tests:
        doc.add_heading("Результат выполнения тестов", level=2)
        doc.add_paragraph(console_test_output)
    doc.save(output_file)
    return output_file


if __name__ == "__main__":
    input_directory = input("Введите путь к лабе: ")
    input_directory = os.path.abspath(input_directory)
    executable = input(
        "Введите название исполняемого файла (оставьте пустым при main.py): "
    )
    report_path = create_docx(input_directory, executable)
    print(f"Отчёт сохранен в: {report_path}")
